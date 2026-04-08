from docx import Document
from docx.shared import Inches
import os

# Define paths
ARTIFACT_DIR = r"C:\Users\DELL\.gemini\antigravity\brain\ca6e7909-9ebd-4469-a14f-04fce9ab4ded"
OUTPUT_FILE = os.path.join(ARTIFACT_DIR, "NutriPet_Opto_EDA_and_Data_Sources.docx")

# Create a new Document
doc = Document()

# Title
doc.add_heading('NutriPet_Opto: Consolidated EDA and Data Sources Report', 0)

doc.add_paragraph("This report consolidates the data sources documentation and the Exploratory Data Analysis (EDA) findings for the NutriPet_Opto project.")

# Part 1
doc.add_heading('Part 1: Data Sources', level=1)
doc.add_paragraph("(Source: DATABASES/DATA_SOURCES.md)")

doc.add_heading('1. Data Sources & Specific Files', level=2)

doc.add_heading('1.1 USDA FoodData Central (Primary Source)', level=3)
doc.add_paragraph("Source URL: https://fdc.nal.usda.gov/download-datasets.html")
doc.add_paragraph("Specific Dataset Used: SR Legacy Foods (Standard Reference)")
p = doc.add_paragraph()
p.add_run("  - File: SR_Legacy_Foods.zip (specifically food.csv and nutrient.csv)\n")
p.add_run("  - Rationale: The Standard Reference dataset provides the most comprehensive historical data on basic commodities (fruits, vegetables, meats) used in pet food, compared to the newer branded food products.")
doc.add_paragraph("Data Extracted:")
p = doc.add_paragraph()
p.add_run("  - Proximates: Protein, Fat, Carbohydrate, Fiber, Sugar, Energy (Kcal)\n")
p.add_run("  - Micronutrients: Fatty acids (Omega-3/6 markers), Calcium, Phosphorus")

doc.add_heading('1.2 ASPCA Animal Poison Control Center', level=3)
doc.add_paragraph("Source URL: https://www.aspca.org/pet-care/animal-poison-control")
doc.add_paragraph("Specific Resources References:")
p = doc.add_paragraph()
p.add_run("  - 'People Foods to Avoid Feeding Your Pets' (Web List)\n")
p.add_run("  - 'Toxic and Non-Toxic Plants List' (Database)")
doc.add_paragraph("Data Extracted:")
p = doc.add_paragraph()
p.add_run("  - List of toxic substances (e.g., Theobromine, Xylitol, Allium compounds)\n")
p.add_run("  - Severity classification (derived from 'clinical signs' descriptions)")

doc.add_heading('1.3 AAFCO & NRC Guidelines', level=3)
doc.add_paragraph("Source: Association of American Feed Control Officials (2023 Official Publication) & NRC Nutrient Requirements of Dogs and Cats.")
doc.add_paragraph("Specific Tables Used:")
p = doc.add_paragraph()
p.add_run("  - 'AAFCO Dog and Cat Food Nutrient Profiles'\n")
p.add_run("  - Used for: Determining protein_requirement, fat_tolerance, and mineral thresholds.")

# Dataset Description
doc.add_heading('2. Dataset Description', level=2)
doc.add_heading('2.1 Overview', level=3)
doc.add_paragraph("This dataset contains detailed nutritional profiles of raw and processed ingredients used in pet food, combined with species-specific nutritional requirements and known toxicity data.")
p = doc.add_paragraph()
p.add_run("- Total records: ~1,056 rows\n")
p.add_run("- Primary format: CSV/TSV -> SQLite database")

doc.add_heading('2.2 Tables & Attributes', level=3)
doc.add_paragraph("The dataset includes tables for Species (8 records), Foods (120+ records), and Toxicity (80+ records).")
doc.add_paragraph("Target Variables:")
p = doc.add_paragraph()
p.add_run("- predicted_caloric_density (regression)\n")
p.add_run("- health_grade (classification)")

# Part 2
doc.add_heading('Part 2: Exploratory Data Analysis (EDA) & Model Performance', level=1)
doc.add_paragraph("(Source: backend/ml/eda_report.md with visualizations)")

doc.add_heading('1. Executive Summary', level=2)
doc.add_paragraph("This report analyzes the synthetic dataset (1,056 records). Key Findings:")
p = doc.add_paragraph()
p.add_run("- Data Balance: Balanced distribution across Health Grades A-F.\n")
p.add_run("- Top Risk Factor: toxin_flag is the dominant predictor for Grade F.\n")
p.add_run("- Model Performance: XGBoost (70.7%) and Random Forest (70.3%) achieved the highest accuracy.")

doc.add_heading('2. Exploratory Data Analysis (EDA)', level=2)

doc.add_heading('2.1 Grade Distribution', level=3)
doc.add_paragraph("The health grade distribution is well-represented across all classes.")
try:
    doc.add_picture(os.path.join(ARTIFACT_DIR, 'grade_distribution.png'), width=Inches(6))
except Exception as e:
    doc.add_paragraph(f"[Error adding grade_distribution.png: {e}]")

doc.add_heading('2.2 Key Correlations', level=3)
doc.add_paragraph("The correlation heatmap reveals strong relationships between toxins, sugar, and health grades.")
try:
    doc.add_picture(os.path.join(ARTIFACT_DIR, 'correlation_heatmap.png'), width=Inches(6))
except Exception as e:
    doc.add_paragraph(f"[Error adding correlation_heatmap.png: {e}]")

doc.add_heading('Specific Feature Correlations', level=4)
doc.add_paragraph("Sugar Impact:")
try:
    doc.add_picture(os.path.join(ARTIFACT_DIR, 'sugar_vs_grade.png'), width=Inches(6))
except Exception as e:
    doc.add_paragraph(f"[Error adding sugar_vs_grade.png: {e}]")

doc.add_paragraph("Protein Impact:")
try:
    doc.add_picture(os.path.join(ARTIFACT_DIR, 'protein_vs_grade.png'), width=Inches(6))
except Exception as e:
    doc.add_paragraph(f"[Error adding protein_vs_grade.png: {e}]")

doc.add_paragraph("Glycemic Impact:")
try:
    doc.add_picture(os.path.join(ARTIFACT_DIR, 'glycemic_vs_grade.png'), width=Inches(6))
except Exception as e:
    doc.add_paragraph(f"[Error adding glycemic_vs_grade.png: {e}]")

doc.add_heading('2.3 Nutrient Analysis', level=3)
doc.add_paragraph("Analysis of toxicity and omega balances.")
try:
    doc.add_picture(os.path.join(ARTIFACT_DIR, 'toxicity_vs_grade.png'), width=Inches(6))
    doc.add_picture(os.path.join(ARTIFACT_DIR, 'omega_balance.png'), width=Inches(6))
    doc.add_picture(os.path.join(ARTIFACT_DIR, 'caloric_density_distribution.png'), width=Inches(6))
except Exception as e:
    doc.add_paragraph(f"[Error adding nutrient analysis plots: {e}]")

doc.add_heading('2.4 Species-Specific Patterns', level=3)
try:
    doc.add_picture(os.path.join(ARTIFACT_DIR, 'species_comparison.png'), width=Inches(6))
except Exception as e:
    doc.add_paragraph(f"[Error adding species_comparison.png: {e}]")

doc.add_heading('3. Model Performance Evaluation', level=2)
doc.add_heading('3.1 Classification Results', level=3)
doc.add_paragraph("Tree-based models (RF, XGB) outperformed Logistic Regression. Accuracy ~70%.")

try:
    doc.add_picture(os.path.join(ARTIFACT_DIR, 'confusion_matrix.png'), width=Inches(6))
except Exception as e:
    doc.add_paragraph(f"[Error adding confusion_matrix.png: {e}]")

doc.add_heading('3.2 Feature Importance', level=3)
doc.add_paragraph("Top predictors: Toxin Flag, Protein Match, Sugar Risk.")

try:
    doc.add_picture(os.path.join(ARTIFACT_DIR, 'feature_importance.png'), width=Inches(6))
    doc.add_picture(os.path.join(ARTIFACT_DIR, 'cost_analysis.png'), width=Inches(6))
except Exception as e:
    doc.add_paragraph(f"[Error adding feature importance plots: {e}]")

doc.add_heading('4. Conclusion', level=2)
doc.add_paragraph("The NutriPet_Opto system successfully bridges the gap between raw data and actionable advice, achieving >70% accuracy with strict safety checks.")

# Save
doc.save(OUTPUT_FILE)
print(f"Document saved to: {OUTPUT_FILE}")
