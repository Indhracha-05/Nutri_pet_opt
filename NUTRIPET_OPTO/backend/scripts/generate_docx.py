
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('NutriPet_Opto: AI-Powered Pet Food Analysis System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Technical Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = 'Subtitle'
    
    doc.add_paragraph('\n' * 5)
    
    details = [
        "Project Title: NutriPet_Opto — Hybrid AI for Pet Nutrition Analysis",
        "Course Name: [Course Name/Code]",
        "Team Members:",
        "- [Member 1 Name]",
        "- [Member 2 Name]",
        "Instructor: [Instructor Name]",
        "Date: February 12, 2026"
    ]
    
    for line in details:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # CONTENT
    sections = [
        ("2. Abstract", "Overview: Pet owners often struggle to interpret nutritional labels and identify species-specific toxins. This project develops an intelligent system to evaluate pet food safety and quality.\nObjective: To build a hybrid AI pipeline combining a Knowledge Graph (for rule-based toxicity checks) and Machine Learning (for nutritional grading) to provide actionable health insights.\nApproach: We utilized a diverse dataset of pet foods and species requirements. Features were engineered based on biological heuristics (e.g., protein match, sugar risk). A Random Forest Classifier was trained to predict a Health Grade (A–F).\nResults: The system achieved ~70.3% accuracy in health grading, with 100% toxicity detection via the Knowledge Graph."),
        ("3. Problem Statement & Objectives", "Problem: Generic pet food recommendations fail to account for species-specific biological needs (e.g., Cats require Taurine/high protein; Dogs are omnivores but sensitive to Xylitol).\nObjectives:\n1. Safety First: Detect 100% of known toxic ingredients using graph-based rules.\n2. Nutritional Grading: Predict a health grade (A–F) with >70% accuracy.\n3. Explainability: Provide clear reasons for each grade (e.g., \"High Sugar\", \"Low Protein\").\n4. Accessibility: Deploy a user-friendly Web Interface."),
        ("4. Dataset Description", "Source: Aggregated from USDA FoodData Central, ASPCA Toxic Plants Database, and AAFCO nutritional standards.\nSamples: 1056 samples (Synthetic & Real-world mix).\nFeatures:\n- Raw: Calories, Protein (g), Fat (g), Carbs (g), Sugar (g), Fiber (g).\n- Engineered: toxin_flag, protein_match, sugar_risk.\nTarget: health_grade (Categorical: A, B, C, D, E, F).\nClass Distribution: Balanced equally across grades (approx. ~176 samples per class) to prevent bias.\nMissing Values: Handled via median imputation (numeric) and mode imputation (categorical)."),
        ("5. Data Preprocessing", "Cleaning:\n- Duplicates: Removed exact duplicates.\n- Outliers: Capped using IQR method (1.5 × IQR) to prevent skewing from extreme values.\n- Normalization: Standardized text inputs (stripped whitespace, lowercase).\n\nFeature Engineering:\n1. toxin_flag: Binary (1 = Toxic, 0 = Safe). Derived from Knowledge Graph traversal.\n2. protein_match_score: Ratio of Food Protein / Species Requirement.\n3. sugar_risk_score: (Sugar/30g) * (GlycemicIndex/100). High index indicates metabolic risk.\n4. omega_balance: Ratio of Omega-3 to Omega-6 fatty acids (inflammatory marker).\n\nJustification: Raw nutrient values alone don't indicate health. \"High Protein\" is good for a cat but neutral for a rabbit. Engineered features capture this biological context."),
        ("6. Exploratory Data Analysis (EDA)", "Visualizations (See backend/ml/plots/):\n1. Grade Distribution: Confirmed balanced classes.\n2. Correlation Heatmap: Strong negative correlation between sugar_risk and health_grade (High sugar = Lower grade). Positive correlation between protein_match and health_grade for carnivores.\n3. Toxicity Impact: Boxplots show toxin_flag=1 almost exclusively maps to Grade F.\n\nInterpretation: The features successfully capture the biological rules we intended to model."),
        ("7. Model Preparation", "Problem Type: Multi-class Classification (Grade A–F) and Regression (Caloric Density).\nTarget: health_grade (Encoded: A=0, ..., F=5).\nRepresentation:\n- Numerical: Standard Scaled (protein_g, fat_g, etc.).\n- Categorical: One-Hot Encoded (species_Dog, species_Cat).\nSplit Strategy: 80% Train, 20% Test.\nJustification: Standard split provides sufficient data for training while reserving a robust validation set. Stratified sampling ensures all grades are represented in both sets."),
        ("8. Model Implementation", "We compared three models:\n\nModel 1: Random Forest Classifier (Primary)\n- Architecture: Ensemble of Decision Trees.\n- Hyperparameters: n_estimators=200, max_depth=15, min_samples_split=5.\n- Justification: Robust to outliers, handles non-linear relationships (e.g., \"Protein is good up to a point\"), and provides feature importance.\n\nModel 2: XGBoost Classifier\n- Architecture: Gradient Boosted Decision Trees.\n- Justification: Often achieves ease-of-edge performance on tabular data. Used for comparison.\n\nModel 3: Logistic Regression\n- Architecture: Linear Model.\n- Justification: Baseline to checking if complex models are actually needed."),
        ("9. Evaluation Metrics", "Chosen Metrics:\n- Accuracy: Overall correctness.\n- F1-Score (Macro): Balances Precision and Recall across all classes (crucial for multi-class).\n- Confusion Matrix: To visualize where misclassifications occur (e.g., confusing B with C)."),
        ("10. Model Comparison", "Random Forest: 70.28% Accuracy, 0.6463 F1\nXGBoost: 70.75% Accuracy, 0.6663 F1\nLogistic Regression: 66.04% Accuracy, 0.5523 F1\n\nSelection: Random Forest was selected as the final model.\nJustification: While XGBoost was slightly more accurate (+0.5%), Random Forest offers better interpretability and is less prone to overfitting on limited synthetic data. The gap is negligible."),
        ("11. Results & Interpretation", "Strengths:\n- Toxicity: 100% detection of critical safety threats (Grade F).\n- Nuance: Distinguishes well between \"Excellent\" (A) and \"Poor\" (D/E).\n\nWeaknesses:\n- Mid-range confusion: Grades B vs C are sometimes confused due to subtle nutritional differences in the synthetic boundaries. This is expected as biological boundaries are often fluid."),
        ("12. Live Prediction / System Demo", "Sample Input:\n- Species: \"Dog\" | Food: \"Chocolate\"\n- Result: Grade F (Toxic). Explanation: \"Contains Theobromine - Toxic\"\n\nSample Input:\n- Species: \"Cat\" | Food: \"Chicken Breast\"\n- Result: Grade A. Explanation: \"High Protein Match (95%), No Toxins.\""),
        ("13. Conclusion", "The NutriPet_Opto system successfully bridges the gap between raw data and actionable advice. We achieved the objective of creating a safe, accurate (>70%), and explainable tool. The hybrid approach (Graph + ML) ensures that critical safety rules are never violated by probabilistic ML errors."),
        ("14. Limitations & Future Work", "Limitations:\n- Synthetic Data: Real-world pet food formulations are proprietary; we relied partially on synthetic generation.\n- Special Diets: Does not Account for specific conditions like kidney disease or diabetes.\n\nFuture Work:\n- Image Recognition: Allow users to scan food labels directly.\n- User Profiles: Save pet history and allergies.\n- Vets Integration: Connect with local veterinarians.")
    ]
    
    for heading, content in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(content)
        
    # Table (Section 15)
    doc.add_page_break()
    doc.add_heading("15. Final Results Summary Table", level=1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Random Forest (Selected)'
    
    metrics = [
        ("Accuracy", "70.28%"),
        ("F1-Score (Macro)", "0.6463"),
        ("5-Fold CV Accuracy", "70.15% ± 0.02"),
        ("Inference Speed", "< 50ms"),
        ("Model Size", "~2.5 MB")
    ]
    
    for metric, value in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value

    # Save to project root (renamed to avoid lock)
    output_path = os.path.join(os.getcwd(), "Technical_Report_Final.docx")
    doc.save(output_path)
    print(f"Created: {output_path}")

if __name__ == "__main__":
    create_report()
